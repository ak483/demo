import os, time, sys, shutil, logging, re, random, getpass
import pandas as pd
from lxml import etree
from My_code.Toolbox.Word_Conversion import Word_Conversion
from docx import Document
# -*- coding: utf-8 -*-
# 输出到文件
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(funcName)s -> %(message)s',
                    filename=rf'.\{os.path.splitext(os.path.split(__file__)[1])[0]}.log')
# 输出到屏幕
console = logging.StreamHandler()
console.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s %(funcName)s -> %(message)s')
console.setFormatter(formatter)
logging.getLogger('').addHandler(console)

sys.path.append(r'D:\untitled1')
from My_code.名动漫.Mdm_API import ArticleClassAPI
from My_code.名动漫.Article_Published.Submit_SQL import Submit_Data
from My_code.名动漫.Article_Published.Get_Data_ import Data_Format

def Description(htmlStr: str) -> str:
    """
    1、内容摘要

    :param htmlStr: 正文内容
    :return: 内容摘要
    """

    # 解析 HTML 代码
    soup = etree.HTML(htmlStr)
    textStr = soup.xpath('//text()')
    # 清洗
    textStr = ''.join(textStr)
    textStr = textStr.replace('\n', '')
    textStr = textStr.replace('\t', '')
    textStr = textStr.replace('\r', '')
    textStr = textStr.replace(' ', '')
    textStr = textStr.replace('　', '')
    textStr = textStr.replace('&nbsp;', '')
    textStr = textStr.replace('&emsp;', '')
    textStr = textStr.replace(' ', '')
    textStr = textStr.replace('\u2003', '')
    # 内容摘要
    description = textStr[:120]
    description+='...'

    return description

def Word_Format(docxName,dayStr,letterStr):
    """
    1、读取文档

    2、转为 HTML

    3、格式化

    :param docxName: 文档名称
    :param dayStr: 日期
    :param letterStr: 用于命名的字母
    :param newsHuabshiBool: True：画师巴士 or False：名动漫官网
    :return: 格式化好的 HTML 代码
    """
    from pydocx import PyDocX
    from lxml import etree

    # 先转换为HTML文件
    html = PyDocX.to_html(FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + '.docx')
    # 替换 <h1> 标签
    html = re.sub('<h\d.?>', '<p><strong>', re.sub('</h\d.?>', '</strong></p>', html))
    with open(fr'{FILE_PATH_DICT["运行文件夹"]}\{docxName}.html', 'w', encoding='utf-8') as f:
        f.write(html)

    # 解析
    html = etree.parse(fr'{FILE_PATH_DICT["运行文件夹"]}\{docxName}.html', etree.HTMLParser())
    # 获取原始字符串
    htmlList = []
    contents = html.xpath('//body/p | //body/li')
    for i in contents:
        # 当前节点字符串
        # contentss=i.xpath('.//text()')
        nodeStr = re.sub('\u3000', '', etree.tostring(i, encoding=str))

        # 防止小图片被添加进去
        for j in i.xpath('.//img'):
            try:
                width = re.search('\d+', str(j.xpath('./@width'))).group()
                height = re.search('\d+', str(j.xpath('./@height'))).group()
                if int(width) <= 100 and int(height) <= 100:
                    # 小图片字符串
                    nodeStr_ = etree.tostring(j, encoding=str)
                    # 去除小图片，保留非图片内容
                    nodeStr__ = re.sub('<img.*/>', '', nodeStr_)
                    # 根据原内容换成新内容
                    nodeStr = nodeStr.replace(nodeStr_, nodeStr__, 1)
                    pass
            except AttributeError:
                pass

        # 获取原始字符串
        htmlList.append(nodeStr)

    # 清洗
    htmlList_ = []
    # 项目编号
    e = 1
    for i in htmlList:
        a = re.sub('</p>', '[/p]', re.sub('<p>', '[p]', i))
        b = re.sub('</strong>', '[/strong]', re.sub('<strong>', '[strong]', a))
        c = re.sub('<img.*?/>', 'img_img', b)
        d = re.sub('<.*?>', '', c)
        # 针对文档使用项目编号这种格式
        if str(i[:4]) == '<li>':
            d = f'{e}.{d}'
            e += 1
            if '[p]' not in d:
                d = f'[p]{d}[/p]'
        else:
            e = 1
        htmlList_.append(d)

    # 格式化
    imgName = 0
    htmlList__ = []
    for i in htmlList_:
        # 转 span
        a = re.sub('\[/p\]', '</p>', re.sub('\[p\]', '<p>', i))
        # 未有加粗
        if not re.search('\[strong\]', a):
            b = a.replace(
                '<p>', '<p>&nbsp; &nbsp; &nbsp; &nbsp; ', 1)
        # 教程该段有无加粗
        else:
            b = re.sub('\[/strong\]', '</strong>', re.sub('\[strong\]', '<strong>', a))
            # 整段加粗
            if '<p><strong>' == b[:35]:
                b = '<br />' + b.replace(
                    '<p>', '<p>&nbsp; &nbsp; &nbsp; &nbsp; ', 1)
            # 单个字加粗
            else:
                b = b.replace(
                    '<p>', '<p>&nbsp; &nbsp; &nbsp; &nbsp; ', 1)
        # 该段有无图片
        if re.search('img_img', b):
            # 去掉空行
            b = re.sub('&nbsp; &nbsp; &nbsp; &nbsp; ', '', b)
            # 检查该段是否有文字
            if len(b) != len('<p>img_img</p>'):
                # 去除图片被加粗
                b = re.sub('<br /><p><strong>img_img</strong>', '<p>img_img', b)
                b = re.sub('<p><strong>img_img</strong>', '<p>img_img', b)
                # 提取非图片的文字内容
                d = re.findall('<.*?>(.*?)<.*?>', re.sub('img_img', '<><>', b))
                # 顺序替换
                for k in d:
                    # 文字内容不为空
                    if k != '' and k != ' ':
                        # 有中文内容
                        if re.search(u'[\u4e00-\u9fa5]', k):
                            # 该内容进行居中
                            b = b.replace(k, '<div style="text-align:center;"><p>' + k + '</p></div><br />', 1)
                        else:
                            b = b.replace(k, '', 1)
            # 顺序更新图片
            for k in range(len(re.findall('img_img', b))):
                # True：画师巴士 or False：名动漫官网

                if len(re.findall('img_img', b)) == 1:
                    b = b.replace('img_img',
                                  f'<div style="text-align:center;">'
                                  f'<img src="{dayStr}/{letterStr}.jpg" alt=""/></div>', 1)
                else:
                    b = b.replace('img_img',
                                  f'<div style="text-align:center;">'
                                  f'<img src="{dayStr}/{letterStr}.jpg" alt=""/></div></br>', 1)
                # 图片个数
                imgName += 1
            htmlList__.append(b)
        # 没有图片
        else:
            if '图片来源于网络' in b:
                htmlList__.append(
                    '<div style="text-align:center;">' + re.sub('&nbsp; &nbsp; &nbsp; &nbsp; ', '', b) + '</div>'
                )
            else:
                htmlList__.append(b)
    # 完整HTML字符串
    htmlStr = '<br />'.join(htmlList__)
    htmlStr = htmlStr.replace('请收藏名动漫官网', '请收藏画师巴士官网')
    htmlStr = htmlStr.replace("'", '"')

    return htmlStr


def Get_Data(mod: str) -> list:
    # 读取模板
    mainExcelDict = pd.read_excel(
        FILE_PATH_DICT['兼职文章1'], sheet_name=[
            'Sheet1',
        ]
    )
    mainExcelData = mainExcelDict['Sheet1']

    # 筛选
    mainExcelData = mainExcelData[(mainExcelData['发布状态'] != '已发布')]
    titleList = mainExcelData['文章标题'].to_list()
    keyList = mainExcelData['关键词'].to_list()
    labelList = mainExcelData['标签'].to_list()
    pictureClassificationList = mainExcelData['配图分类'].to_list()
    # 自定义课程
    customizeCourseList = mainExcelData['推荐课程ID'].to_list()
    articleClassificationNameList = mainExcelData['文章分类（名动漫）'].to_list()
    MDM_article = mainExcelData['文章存储路径'].to_list()

    dataSetList = []
    # 自动删除文档
    removeFileList = []
    for item in range(len(MDM_article)):
        if len(dataSetList) >= MAXINDEX:
            break

        # 配图1，配图2
        if pictureClassificationList[item] == '3D':
            a, b = random.sample([i for i in range(0, 100)], 2)
        elif pictureClassificationList[item] == '原画':
            a, b = random.sample([i for i in range(0, 142)], 2)
        elif pictureClassificationList[item] == '影视':
            a, b = random.sample([i for i in range(0, 100)], 2)
        elif pictureClassificationList[item] == '插画':
            a, b = random.sample([i for i in range(0, 112)], 2)
        elif pictureClassificationList[item] == '游戏UI':
            a, b = random.sample([i for i in range(0, 100)], 2)
        elif pictureClassificationList[item] == '漫画':
            a, b = random.sample([i for i in range(0, 100)], 2)
        elif pictureClassificationList[item] == '素描':
            a, b = random.sample([i for i in range(0, 118)], 2)
        elif pictureClassificationList[item] == '线稿':
            a, b = random.sample([i for i in range(0, 181)], 2)
        elif pictureClassificationList[item] == '绘画':
            a, b = random.sample([i for i in range(0, 172)], 2)
        else:
            a, b = random.sample([i for i in range(0, 100)], 2)

            # 检查配图分类正不正确
        try:
            FILE_PATH_DICT['配图'][pictureClassificationList[item]]
        except KeyError:
            logging.info(f'配图分类不存在：{pictureClassificationList[item]}')
            continue


        contentImageUrlList = [
            rf'<div style="text-align:center;"><img src="{FILE_PATH_DICT["配图"][pictureClassificationList[item]]}/{a}.jpg" alt="" /></div><br />',
            rf'<div style="text-align:center;"><img src="{FILE_PATH_DICT["配图"][pictureClassificationList[item]]}/{b}.jpg" alt="" /></div><br />',
        ]
        # 封面图
        coverImageUrl = rf"{FILE_PATH_DICT['配图封面'][pictureClassificationList[item]]}/{a}.jpg"
        timeDateList = ''
        try:
            timeDate = int(time.mktime(time.strptime(str(timeDateList), '%Y/%m/%d %H:%M'))) if timeDateList != '' else ''
        except Exception:
            timeDate = ''

        #读取文档内容
        while True:
            # 文档名称
            docxName = ''
            # 文件路径
            if os.path.isfile(MDM_article[item]):
                # 文档路径，文档名称与后缀
                docxPath, docxNameSuf = os.path.split(MDM_article[item])
                # 文档名称，文档后缀
                docxName, docxSuf = os.path.splitext(docxNameSuf)
                # 复制到本地
                shutil.copyfile(MDM_article[item], FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + docxSuf)



                if '.doc' == docxSuf:
                    msg = Word_Conversion(
                        docFilePath=FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + docxSuf,
                        docxFileSavePath=FILE_PATH_DICT['运行文件夹']
                    )
                    if isinstance(msg, str):
                        logging.info(f'doc文档转换为docx失败 -> {msg}')

                    # 插入图片
                    doc = Document(FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + '.docx')
                    doc.paragraphs[3].text = 'img_img'
                    doc.save(FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + '.docx')

                else:

                    doc = Document(FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + '.docx')
                    doc.paragraphs[3].text = 'img_img'
                    doc.save(FILE_PATH_DICT['运行文件夹'] + '\\' + docxName + '.docx')


            else:

                logging.info('文件或者文件夹路径不存在！！！')
                logging.info(MDM_article[item])
                logging.info('输入新的文件或者文件夹路径：')
                MDM_article[item] = str(input())
                continue
            break

        if docxName == '':
            logging.info('未能找到文档！！！')
            logging.info(MDM_article[item])
            continue

        # 添加文档到删除列表
        removeFileList.append(docxName)
        htmlStr = Word_Format(docxName,FILE_PATH_DICT["配图"][pictureClassificationList[item]],a)

        dataSetList.append({
            # 文章标题
            'articleTitle': titleList[item],
            # 短标题
            'articleShortTitle': re.sub('(\(.*?\))', '', titleList[item]),
            # 短标题未去除末尾符号（用于双1去重）
            'shortTitle': re.sub('(\(.*?\))', '', titleList[item]),
            # 文章分类
            'articleClassificationName': articleClassificationNameList[item],
            # 关键词
            'keywordsStr': keyList[item],
            # 标签
            'labelStr': labelList[item],
            # 正文内容
            'content': htmlStr,
            # # 核心内容
            # 'content_': content,
            # 内容摘要
            'descriptionStr': Description(htmlStr),
            # 封面图
            'coverImageUrl': coverImageUrl,
            # 配图分类
            'pictureClassification': pictureClassificationList[item],
            # # 首段模板
            # 'firstParagraph': firstParagraph_,
            # # 中段模板
            # 'middleSection': re.sub(
            #     '<p>&nbsp; &nbsp; &nbsp; &nbsp; ', '', re.sub(
            #         '</p><br />', '', re.sub(newTitle, '{标题}', middleSection))),
            # # 尾段模板
            # 'tailSection': tailSection_,
            # # 名动漫官网课程ID
            'mdmCourseIdInt': customizeCourseList[item],
            # 发布日期
            'timeDate': timeDate,
            # 'articleIdInt': int(item+29646),
        })

    return dataSetList





def Key_Html(articleId: int, keyStr: str, htmlStr: str):
    """
    1、关键词添加链接

    :param articleId: 添加成功的 ID
    :param keyStr: 关键词
    :param htmlStr: 格式化好的 HTML 代码
    :return: 添加好的 HTML 代码
    """

    htmlEtree = etree.HTML(htmlStr)
    htmlEtreeText = ''.join(htmlEtree.xpath('//text()'))
    # 关键词添加链接
    for key in keyStr.split(','):
        if key in htmlEtreeText:
            try:
                # 关键词有无添加链接
                keyBool = re.search(f'html">{key}</a>', htmlStr).group()
            except AttributeError:
                keyBool = False
            # 没有添加链接则添加
            if keyBool is False:
                htmlStr = htmlStr.replace(key, f'<a href="https://www.mingdongman.com/news/{int(articleId)}.html">{key}</a>', 1)

    return htmlStr


def Mdm_Article_Run(articleDataList: list):
    """
    1、发布半自动洗稿项目

    :param articleDataList:
    :return:
    """

    if isinstance(articleDataList, bool):
        return True
    if len(articleDataList) == 0:
        return True

    # 实例化API
    articleClassAPI = ArticleClassAPI()

    # 标题
    excelTitleList = []
    # 发布文章
    for item in articleDataList:
        msg = articleClassAPI.Add_Article(
            articleTitle=item['articleTitle'], classify=item['articleClassificationName'], readingVolumeInt=random.randint(1, 50),
            originalBool=1,  releaseTimeInt=item['timeDate'] if str(item['timeDate']) != '' else int(str(time.time())[:10]),
            labelList=item['labelStr'].split(','), keywordsList=item['labelStr'].split(','),
            thumbnailStr=item['coverImageUrl'], descriptionStr=item['descriptionStr'], contentStr=item['content'],
            shortTitle=item['articleShortTitle'], newsHuabshiBool=False, courseIdInt=item['mdmCourseIdInt']
        )
        if isinstance(msg, list):
            # 文章ID
            articleId = int(msg[0]['idInt'])
            logging.info(f'添加文章成功：{articleId}')
            # 添加关键词链接`
            htmlStr = Key_Html(articleId=articleId, keyStr=item['labelStr'], htmlStr=item['content'])
            # 更新内容
            msg = articleClassAPI.Update_Article(articleIdInt=articleId, contentStr=htmlStr,  newsHuabshiBool=False)
            if isinstance(msg, list):
                pass
            else:
                logging.info(f'更新该文章的内容失败：{articleId} -> 原因：{msg}')

            excelTitleList.append(item['articleTitle'])
            (pd.DataFrame({'名动漫官网标题': excelTitleList})).to_excel(FILE_PATH_DICT['标题保存'], index=False)

            # # 存档
            # Submit_Data(
            #     articleid=str(articleId), platform='名动漫官网', title=item['articleTitle'], shorttitle=item['shortTitle'],
            #     classification=item['articleClassificationName'], keywords=item['labelStr'], coverimageurl=item['coverImageUrl'],
            #     pictureclassification=item['pictureClassification'], firstparagraph=item['firstParagraph'], middlesection=item['middleSection'],
            #     tailsection=item['tailSection'], content=item['content_']
            # )
        else:
            logging.info(f"添加该文章失败：{item['articleTitle']} -> 原因：{msg}")
            excelTitleList.append('发布失败')
            (pd.DataFrame({'名动漫官网标题': excelTitleList})).to_excel(FILE_PATH_DICT['标题保存'], index=False)

        time.sleep(1)

    pass


FILE_PATH_DICT = {
    '半自动洗稿项目': r'\\DESKTOP-J6ECV53\Users\Adminitrator03\Desktop\脚本运行产生的文件\半自动洗稿\半自动洗稿项目-名动漫官网.xlsx',
    '兼职文章': r'\\DESKTOP-J6ECV53\Users\Adminitrator03\Desktop\兼职文章.xlsx',
    '兼职文章1': r'\\DESKTOP-J6ECV53\Users\Adminitrator03\Desktop\兼职文章1.xlsx',
    '标题保存': r'\\DESKTOP-J6ECV53\Users\Adminitrator03\Desktop\脚本运行产生的文件\半自动洗稿\标题2.xlsx',
    '原文配图': 'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/Fifth_Batch',
    '运行文件夹': fr'C:\Users\{getpass.getuser()}\Desktop\脚本运行产生的文件',
    '配图': {
        '3D': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/3D',
        '插画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E6%8F%92%E7%94%BB',
        '绘画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E7%BB%98%E7%94%BB',
        '漫画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E6%BC%AB%E7%94%BB',
        '素描': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E7%B4%A0%E6%8F%8F',
        '线稿': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E7%BA%BF%E7%A8%BF',
        '影视': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E5%BD%B1%E8%A7%86',
        '游戏UI': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E6%B8%B8%E6%88%8FUI',
        '原画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/%E5%8E%9F%E7%94%BB',
    },
    '配图封面': {
        '3D': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/3D',
        '插画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E6%8F%92%E7%94%BB',
        '绘画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E7%BB%98%E7%94%BB',
        '漫画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E6%BC%AB%E7%94%BB',
        '素描': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E7%B4%A0%E6%8F%8F',
        '线稿': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E7%BA%BF%E7%A8%BF',
        '影视': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E5%BD%B1%E8%A7%86',
        '游戏UI': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E6%B8%B8%E6%88%8FUI',
        '原画': r'https://mdm-article.oss-cn-shenzhen.aliyuncs.com/Article_Image/coverImage/%E5%8E%9F%E7%94%BB',
    },

}

if __name__ == '__main__':
    # 发布文章数目
    MAXINDEX = 30


    Mdm_Article_Run(Get_Data(mod='名动漫小站-名动漫'))
    pass
